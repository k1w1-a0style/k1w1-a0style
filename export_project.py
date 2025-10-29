#!/usr/bin/env python3
# export_project.py — robustes Projekt-Export-Script
import os
import re
from docx import Document

# === KONFIG ===
project_path = '/home/k1w1/k1w1-a0style'
output_txt = '/home/k1w1/output.txt'
output_docx = '/home/k1w1/output.docx'
included_extensions = ['.js', '.css', '.html', '.json', '.md', '.ts', '.tsx', '.jsx', '.config.js']
# Pfade/Ordner, die auf jeden Fall übersprungen werden
skip_dirs = {'node_modules', '.git', 'android', 'ios', 'build', 'dist', '__pycache__'}
# Max. Datei-Größe (bytes) die komplett eingelesen wird — größere Dateien werden als "excerpt" behandelt
max_file_size = 1 * 1024 * 1024  # 1 MB
# Max. Zeilen pro Datei im DOCX (schützt vor riesigen Einträgen)
max_lines_per_file = 1000

# Regex: entferne Steuerzeichen außer Tab(09), LF(10), CR(13)
_ctrl_re = re.compile(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]')

def is_text_bytes(b, threshold=0.30):
    """Heuristik: Anteil nicht-text Bytes (NULL/control) -> wenn zu hoch => binär."""
    if not b:
        return True
    # Count control bytes
    control = sum(1 for ch in b if ch < 32 and ch not in (9,10,13))
    return (control / len(b)) < threshold

def clean_bytes(b):
    """Entfernt NUL-Bytes und andere problematische Steuerzeichen."""
    # ersetze NUL durch nichts, dann decode mit replace
    b = b.replace(b'\x00', b'')
    try:
        s = b.decode('utf-8')
    except Exception:
        s = b.decode('utf-8', errors='replace')
    # entferne restliche Steuerzeichen
    s = _ctrl_re.sub('', s)
    return s

def should_include(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in included_extensions:
        return True
    return False

def export_to_txt(project_path, and_file, included_extensions):
    with open(and_file, 'w', encoding='utf-8') as outfile:
        for foldername, dirnames, filenames in os.walk(project_path):
            # skip dirs
            dirnames[:] = [d for d in dirnames if d not in skip_dirs]
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                if not should_include(file_path):
                    continue
                try:
                    size = os.path.getsize(file_path)
                    mode = 'rb'
                    with open(file_path, mode) as infile:
                        if size > max_file_size:
                            # read head only
                            head = infile.read(200*1024)
                            text = clean_bytes(head)
                            outfile.write(f"{'-'*80}\n")
                            outfile.write(f"Datei (excerpt, {size} bytes): {file_path}\n")
                            outfile.write(f"{'-'*80}\n")
                            outfile.write(text + "\n\n")
                            outfile.write(f"[... Datei zu groß, nur Auszug exportiert ...]\n\n")
                        else:
                            b = infile.read()
                            if not is_text_bytes(b):
                                outfile.write(f"{'-'*80}\n")
                                outfile.write(f"Datei (vermutlich binär, übersprungen): {file_path}\n")
                                outfile.write(f"{'-'*80}\n\n")
                                continue
                            text = clean_bytes(b)
                            outfile.write(f"{'-'*80}\n")
                            outfile.write(f"Datei: {file_path}\n")
                            outfile.write(f"{'-'*80}\n")
                            outfile.write(text + "\n\n")
                except Exception as e:
                    outfile.write(f"[Fehler beim Lesen der Datei {file_path}: {e}]\n\n")

def export_to_docx(project_path, output_docx, included_extensions):
    doc = Document()
    for foldername, dirnames, filenames in os.walk(project_path):
        dirnames[:] = [d for d in dirnames if d not in skip_dirs]
        for filename in filenames:
            file_path = os.path.join(foldername, filename)
            if not should_include(file_path):
                continue
            try:
                size = os.path.getsize(file_path)
                with open(file_path, 'rb') as infile:
                    if size > max_file_size:
                        head = infile.read(200*1024)
                        text = clean_bytes(head)
                        doc.add_heading(f"Datei (excerpt, {size} bytes): {file_path}", level=2)
                        # nur erste X Zeilen einfügen
                        lines = text.splitlines()[:max_lines_per_file]
                        for ln in lines:
                            doc.add_paragraph(ln)
                        doc.add_paragraph("[... Datei zu groß, nur Auszug exportiert ...]")
                        doc.add_page_break()
                        continue
                    b = infile.read()
                    if not is_text_bytes(b):
                        doc.add_heading(f"Datei (vermutlich binär, übersprungen): {file_path}", level=2)
                        doc.add_paragraph("[Binärdatei: übersprungen]")
                        doc.add_page_break()
                        continue
                    text = clean_bytes(b)
                    doc.add_heading(f"Datei: {file_path}", level=1)
                    # split in Zeilen, begrenzen
                    lines = text.splitlines()[:max_lines_per_file]
                    for ln in lines:
                        doc.add_paragraph(ln)
                    if len(text.splitlines()) > max_lines_per_file:
                        doc.add_paragraph("[... Ausgabe gekürzt ...]")
                    doc.add_page_break()
            except Exception as e:
                doc.add_paragraph(f"[Fehler bei Datei {file_path}: {e}]")
                doc.add_page_break()
    doc.save(output_docx)

if __name__ == '__main__':
    print("Starte Export...")
    export_to_txt(project_path, output_txt, included_extensions)
    export_to_docx(project_path, output_docx, included_extensions)
    print("Export abgeschlossen:")
    print(f"- TXT: {output_txt}")
    print(f"- DOCX: {output_docx}")
